"""Excel/Word hisobot generatori — database.get_report_data() natijasidan fayl yasaydi."""

import io

from docx import Document
from openpyxl import Workbook
from openpyxl.styles import Font

PERIOD_LABEL = {"week": "Haftalik", "month": "Oylik"}
INTENT_LABEL = {
    "question": "Savol",
    "order": "Buyurtma",
    "feedback": "Fikr",
    "complaint": "Shikoyat",
    "other": "Boshqa",
}
SENTIMENT_LABEL = {"positive": "Ijobiy", "neutral": "Neytral", "negative": "Salbiy"}


def _avg_response_text(data: dict) -> str:
    return str(data["avgResponseMinutes"]) if data["avgResponseMinutes"] is not None else "—"


def build_xlsx(data: dict, business_name: str) -> bytes:
    wb = Workbook()
    ws = wb.active
    ws.title = "Hisobot"
    bold = Font(bold=True)

    ws.append([f"{PERIOD_LABEL[data['period']]} hisobot — {business_name}"])
    ws["A1"].font = Font(bold=True, size=14)
    ws.append([f"Davr: {data['since']} — {data['until']}"])
    ws.append([])

    ws.append(["Ko'rsatkich", "Qiymat"])
    for cell in ws[ws.max_row]:
        cell.font = bold
    ws.append(["Mijoz xabarlari", data["totalCustomerMessages"]])
    ws.append(["Biznes javoblari", data["totalBusinessReplies"]])
    ws.append(["Faol suhbatlar", data["uniqueConversations"]])
    ws.append(["O'rtacha javob vaqti (daqiqa)", _avg_response_text(data)])
    ws.append([])

    ws.append(["Kayfiyat", "Soni"])
    for cell in ws[ws.max_row]:
        cell.font = bold
    for key, label in SENTIMENT_LABEL.items():
        ws.append([label, data["sentiment"].get(key, 0)])
    ws.append([])

    ws.append(["Murojaat turi", "Soni"])
    for cell in ws[ws.max_row]:
        cell.font = bold
    for item in data["intents"]:
        ws.append([INTENT_LABEL.get(item["intent"], item["intent"]), item["count"]])
    ws.append([])

    ws.append(["Sana", "Mijoz xabarlari", "Biznes javoblari"])
    for cell in ws[ws.max_row]:
        cell.font = bold
    for day in data["daily"]:
        ws.append([day["date"], day["customerMessages"], day["businessReplies"]])

    for col, width in zip("ABC", (30, 20, 20)):
        ws.column_dimensions[col].width = width

    buffer = io.BytesIO()
    wb.save(buffer)
    return buffer.getvalue()


def build_docx(data: dict, business_name: str) -> bytes:
    doc = Document()
    doc.add_heading(f"{PERIOD_LABEL[data['period']]} hisobot — {business_name}", level=1)
    doc.add_paragraph(f"Davr: {data['since']} — {data['until']}")

    doc.add_heading("Umumiy ko'rsatkichlar", level=2)
    table = doc.add_table(rows=1, cols=2)
    table.style = "Light Grid Accent 1"
    table.rows[0].cells[0].text = "Ko'rsatkich"
    table.rows[0].cells[1].text = "Qiymat"
    for label, value in [
        ("Mijoz xabarlari", str(data["totalCustomerMessages"])),
        ("Biznes javoblari", str(data["totalBusinessReplies"])),
        ("Faol suhbatlar", str(data["uniqueConversations"])),
        ("O'rtacha javob vaqti (daqiqa)", _avg_response_text(data)),
    ]:
        cells = table.add_row().cells
        cells[0].text, cells[1].text = label, value

    doc.add_heading("Mijozlar kayfiyati", level=2)
    table2 = doc.add_table(rows=1, cols=2)
    table2.style = "Light Grid Accent 1"
    table2.rows[0].cells[0].text = "Kayfiyat"
    table2.rows[0].cells[1].text = "Soni"
    for key, label in SENTIMENT_LABEL.items():
        cells = table2.add_row().cells
        cells[0].text, cells[1].text = label, str(data["sentiment"].get(key, 0))

    doc.add_heading("Murojaat turlari", level=2)
    table3 = doc.add_table(rows=1, cols=2)
    table3.style = "Light Grid Accent 1"
    table3.rows[0].cells[0].text = "Tur"
    table3.rows[0].cells[1].text = "Soni"
    for item in data["intents"]:
        cells = table3.add_row().cells
        cells[0].text = INTENT_LABEL.get(item["intent"], item["intent"])
        cells[1].text = str(item["count"])

    doc.add_heading("Kunlik dinamika", level=2)
    table4 = doc.add_table(rows=1, cols=3)
    table4.style = "Light Grid Accent 1"
    table4.rows[0].cells[0].text = "Sana"
    table4.rows[0].cells[1].text = "Mijoz xabarlari"
    table4.rows[0].cells[2].text = "Biznes javoblari"
    for day in data["daily"]:
        cells = table4.add_row().cells
        cells[0].text = day["date"]
        cells[1].text = str(day["customerMessages"])
        cells[2].text = str(day["businessReplies"])

    buffer = io.BytesIO()
    doc.save(buffer)
    return buffer.getvalue()
